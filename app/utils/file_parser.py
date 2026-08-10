"""
文件解析器 — 支持多格式文件内容提取

支持格式:
  - .txt  : 纯文本
  - .md   : Markdown（去除格式标记）
  - .pdf  : PDF（使用 pdfminer）
  - .docx : Word 文档（使用 python-docx）
"""

import os
import re
from typing import Tuple

from app.core.logger import logger


def extract_text(file_path: str) -> Tuple[str, str]:
    """
    根据文件扩展名自动提取文本内容。

    Args:
        file_path: 文件路径

    Returns:
        (文本内容, 文件名)
    """
    filename = os.path.basename(file_path)
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".txt":
        return _extract_txt(file_path), filename
    elif ext == ".md":
        return _extract_md(file_path), filename
    elif ext == ".pdf":
        return _extract_pdf(file_path), filename
    elif ext == ".docx":
        return _extract_docx(file_path), filename
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """
    从内存中的文件内容提取文本（用于上传接口）。

    Args:
        content: 文件二进制内容
        filename: 文件名（用于判断扩展名）

    Returns:
        提取的文本内容
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".txt":
        return content.decode("utf-8", errors="replace")
    elif ext == ".md":
        raw = content.decode("utf-8", errors="replace")
        return _strip_markdown(raw)
    elif ext == ".pdf":
        return _extract_pdf_from_bytes(content)
    elif ext == ".docx":
        return _extract_docx_from_bytes(content)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


# ── TXT ────────────────────────────────────────────────────────────────────

def _extract_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        return f.read()


# ── Markdown ───────────────────────────────────────────────────────────────

def _strip_markdown(text: str) -> str:
    """去除 Markdown 格式标记，保留纯文本内容"""
    # 移除代码块 (``` ... ```)
    text = re.sub(r"```[\s\S]*?```", "", text)
    # 移除行内代码 (`code`)
    text = re.sub(r"`[^`]+`", "", text)
    # 移除图片 ![alt](url)
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    # 移除链接 [text](url)，保留 text
    text = re.sub(r"\[([^\]]+)\]\(.*?\)", r"\1", text)
    # 移除标题标记 #
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    # 移除粗体/斜体标记
    text = re.sub(r"\*{1,3}([^*]+)\*{1,3}", r"\1", text)
    text = re.sub(r"_{1,3}([^_]+)_{1,3}", r"\1", text)
    # 移除分隔线
    text = re.sub(r"^[-*_]{3,}\s*$", "", text, flags=re.MULTILINE)
    # 移除列表标记
    text = re.sub(r"^[\s]*[-*+]\s+", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*\d+\.\s+", "", text, flags=re.MULTILINE)
    # 移除引用标记
    text = re.sub(r"^>\s?", "", text, flags=re.MULTILINE)
    # 合并空白行
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _extract_md(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        raw = f.read()
    return _strip_markdown(raw)


# ── PDF ────────────────────────────────────────────────────────────────────

def _extract_pdf(file_path: str) -> str:
    """使用 pdfminer 提取 PDF 文本"""
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(file_path)
        return text.strip() or "【PDF 未能提取到文本内容】"
    except Exception as e:
        logger.error(f"[Parser] PDF 解析失败: {e}")
        return ""


def _extract_pdf_from_bytes(content: bytes) -> str:
    """从字节流中提取 PDF 文本"""
    try:
        from io import BytesIO
        from pdfminer.high_level import extract_text as pdfminer_extract
        text = pdfminer_extract(BytesIO(content))
        return text.strip() or "【PDF 未能提取到文本内容】"
    except Exception as e:
        logger.error(f"[Parser] PDF 内存解析失败: {e}")
        return ""


# ── DOCX ───────────────────────────────────────────────────────────────────

def _extract_docx(file_path: str) -> str:
    """使用 python-docx 提取 Word 文档文本"""
    try:
        from docx import Document
        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs) if paragraphs else "【DOCX 未能提取到文本内容】"
    except Exception as e:
        logger.error(f"[Parser] DOCX 解析失败: {e}")
        return ""


def _extract_docx_from_bytes(content: bytes) -> str:
    """从字节流中提取 Word 文档文本"""
    try:
        from io import BytesIO
        from docx import Document
        doc = Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs) if paragraphs else "【DOCX 未能提取到文本内容】"
    except Exception as e:
        logger.error(f"[Parser] DOCX 内存解析失败: {e}")
        return ""


# ── 命令行测试 ─────────────────────────────────────────────────────────────

if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("用法: python file_parser.py <文件路径>")
        sys.exit(1)
    text, name = extract_text(sys.argv[1])
    print(f"文件: {name}")
    print(f"长度: {len(text)} 字符")
    print("─" * 40)
    print(text[:500])